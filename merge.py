import os
from docx import Document


def merge_docx_files(input_directory, output_docx_path):
    """
    Merge multiple DOCX files into a single DOCX file.
    """
    merged_doc = Document()  # Create a new Document for merging

    # Loop through all files in the directory
    for filename in os.listdir(input_directory):
        file_path = os.path.join(input_directory, filename)

        if filename.endswith(".docx"):  # Check if it's a .docx file
            try:
                doc = Document(file_path)  # Load the .docx file

                # Append the content from the current doc to the merged doc
                for para in doc.paragraphs:
                    merged_doc.add_paragraph(para.text)

                print(f"Merged {filename}")
            except Exception as e:
                print(f"Error merging {filename}: {str(e)}")

    # Save the merged document to the output path
    merged_doc.save(output_docx_path)
    print(f"Successfully saved merged DOCX to {output_docx_path}")


def merge_txt_files(input_directory, output_txt_path):
    """
    Merge multiple TXT files into a single TXT file.
    """
    with open(output_txt_path, 'w', encoding='utf-8') as output_file:
        # Loop through all files in the directory
        for filename in os.listdir(input_directory):
            file_path = os.path.join(input_directory, filename)

            if filename.endswith(".txt"):  # Check if it's a .txt file
                try:
                    with open(file_path, 'r', encoding='utf-8') as input_file:
                        output_file.write(input_file.read())  # Append the content from the current file
                        output_file.write("\n\n")  # Add space between files
                    print(f"Merged {filename}")
                except Exception as e:
                    print(f"Error merging {filename}: {str(e)}")

    print(f"Successfully saved merged TXT to {output_txt_path}")


# Example usage
input_directory = 'data/'  # Specify the directory containing the DOCX and TXT files
output_docx_path = 'merged_output.docx'  # Specify the output path for the merged DOCX
# output_txt_path = 'merged_output.txt'  # Specify the output path for the merged TXT

# Call functions to merge DOCX and TXT files
merge_docx_files(input_directory, output_docx_path)
# merge_txt_files(input_directory, output_txt_path)
